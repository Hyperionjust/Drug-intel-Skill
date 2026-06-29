#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_report_docx.py  --  Word (.docx) report generator for the `drug-intel` skill.

    python scripts/build_report_docx.py <data.json> -o <output.docx>

Reads the SAME structured JSON as build_report.py (schema: templates/report_schema.json)
and produces a polished, editable Word document with the identical content & badges:
identity, companies, R&D/approval timeline (table), indication matrix, sales table,
LOE / generics / peers / advantages, the red "需人工核对清单", and a final reference list.

Same rules as SKILL.md: every value keeps its badge(s); red stays loud; checklist always
present; all source links collected into the final 参考来源 (Reference list) section.

Requires python-docx (`pip install python-docx`).
"""

import sys
import json
import argparse

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BADGE_LABEL = {
    "official": "\U0001F3DB官方", "db": "✅库", "web": "\U0001F50Dweb",
    "weak": "⚠存疑", "redflag": "\U0001F534必核", "na": "❓N/A",
}
BADGE_HEX = {
    "official": "1D4ED8", "db": "15803D", "web": "0E7490",
    "weak": "B45309", "redflag": "DC2626", "na": "6B7280",
}
ALIASES = {"fda": "official", "gov": "official", "company": "official", "regulator": "official",
           "database": "db", "pubchem": "db", "chembl": "db", "ctgov": "db", "openfda": "db",
           "media": "web", "secondary": "weak", "unknown": "na", "notfound": "na", "red": "redflag"}


def _norm(k):
    if not k:
        return None
    k = str(k).strip()
    return k if k in BADGE_LABEL else ALIASES.get(k.lower())


def _badge_keys(badge):
    if badge is None:
        return []
    if isinstance(badge, (list, tuple)):
        return [b for b in (_norm(x) for x in badge) if b]
    n = _norm(badge)
    return [n] if n else []


def add_badges(par, badge):
    for k in _badge_keys(badge):
        r = par.add_run("  " + BADGE_LABEL[k])
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(BADGE_HEX[k])


def shade(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexfill)
    tcPr.append(shd)


def set_widths(table, widths_in):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_in):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)


def add_hyperlink(par, url, text, color="1D4ED8"):
    part = par.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    col = OxmlElement("w:color")
    col.set(qn("w:val"), color)
    rPr.append(col)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    par._p.append(link)


def is_us(region):
    r = region or ""
    return r in ("US", "us", "美国") or r.lower().startswith("us")


def date_key(d):
    try:
        return [int(x) for x in str(d).split("-")]
    except Exception:
        return [0]


def banner(doc, text, hexfill, hextext):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(hextext)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexfill)
    pPr.append(shd)
    return p


def kv_section(doc, title, items):
    if not items:
        return
    doc.add_heading(title, level=1)
    for it in items:
        if isinstance(it, str):
            doc.add_paragraph(it, style="List Bullet")
            continue
        p = doc.add_paragraph(style="List Bullet")
        label = it.get("label") or it.get("field")
        text = it.get("text") or it.get("value") or ""
        if label:
            rb = p.add_run(str(label) + "：")
            rb.bold = True
        p.add_run(str(text))
        add_badges(p, it.get("badge"))


def header_table(doc, rows, widths, header_fill="EEF2F7", first_col_fill=None):
    t = doc.add_table(rows=1, cols=len(rows[0]))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(rows[0]):
        hdr[i].text = ""
        rp = hdr[i].paragraphs[0].add_run(htext)
        rp.bold = True
        shade(hdr[i], header_fill)
    for row in rows[1:]:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
        if first_col_fill:
            shade(cells[0], first_col_fill)
    set_widths(t, widths)
    return t


def main(argv):
    ap = argparse.ArgumentParser(description="Render a drug-intel Word (.docx) report from JSON.")
    ap.add_argument("data")
    ap.add_argument("-o", "--out", default=None)
    args = ap.parse_args(argv)
    data = json.load(open(args.data, encoding="utf-8"))
    meta = data.get("meta", {})

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    # Title
    inn = meta.get("inn") or meta.get("name") or "药物情报报告"
    doc.add_heading("药物情报报告：" + inn, level=0)
    sub_bits = []
    for k, pre in (("cn_name", ""), ("brand_us", "\U0001F1FA\U0001F1F8 "), ("brand_cn", "\U0001F1E8\U0001F1F3 ")):
        if meta.get(k):
            sub_bits.append(pre + str(meta[k]))
    if meta.get("codes"):
        sub_bits.append(" / ".join(meta["codes"]))
    if sub_bits:
        doc.add_paragraph("  ·  ".join(sub_bits))
    metabits = ["美国 \U0001F1FA\U0001F1F8 + 中国 \U0001F1E8\U0001F1F3", "生成 " + str(meta.get("generated", ""))]
    if meta.get("mode"):
        metabits.append("口径：" + str(meta["mode"]))
    mp = doc.add_paragraph("  ·  ".join(metabits))
    mp.runs[0].font.size = Pt(9)
    mp.runs[0].font.color.rgb = RGBColor.from_string("6B7280")

    banner(doc, "⚠ 仅供研发情报参考，非医疗建议、非投资建议。所有标 \U0001F534 的字段必须人工复核后再使用。",
           "FFFBEB", "92400E")
    if meta.get("skipped_gates"):
        banner(doc, "⚠ 已按用户要求一次跑完，未逐项停下确认；下面所有标 \U0001F534 的地方都还没经过人工核对。",
               "FEF2F2", "991B1B")

    # legend
    doc.add_heading("徽标说明", level=1)
    lp = doc.add_paragraph()
    for k in ("official", "db", "web", "weak", "redflag", "na"):
        r = lp.add_run(BADGE_LABEL[k] + "  ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(BADGE_HEX[k])
    doc.add_paragraph("🏛️官方=政府/监管/企业官网的客观事实(最高可信) ｜ ✅库=数据库可核 ｜ 🔍web=联网附链接 ｜ "
                      "⚠存疑 ｜ 🔴必核=高幻觉须人工核 ｜ ❓N/A=查过未找到。销售额与主观对比即使来自官方也仍 🔴。")

    # identity
    ident = data.get("identity") or []
    if ident:
        doc.add_heading("身份卡", level=1)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Table Grid"
        for it in ident:
            cells = t.add_row().cells
            cells[0].text = ""
            rr = cells[0].paragraphs[0].add_run(str(it.get("field", "")))
            rr.bold = True
            shade(cells[0], "F1F5F9")
            cells[1].text = str(it.get("value", ""))
            add_badges(cells[1].paragraphs[0], it.get("badge"))
        set_widths(t, [2.0, 4.5])

    # companies
    comp = data.get("companies")
    if comp:
        doc.add_heading("厂家（原研 / 权利人 / 上市持有人）", level=1)
        for flag, label, key, fill in (("\U0001F1FA\U0001F1F8", "美国", "us", "DBEAFE"),
                                       ("\U0001F1E8\U0001F1F3", "中国", "china", "FEE2E2")):
            c = comp.get(key)
            if not c:
                continue
            hp = doc.add_paragraph()
            hr = hp.add_run(flag + " " + label)
            hr.bold = True
            for fld, fl in (("originator", "原研"), ("holder", "当前权利人/上市持有人")):
                v = c.get(fld)
                val = v.get("value") if isinstance(v, dict) else v
                bd = v.get("badge") if isinstance(v, dict) else c.get(fld + "_badge")
                p = doc.add_paragraph(style="List Bullet")
                rb = p.add_run(fl + "：")
                rb.bold = True
                p.add_run(str(val or ""))
                add_badges(p, bd)

    kv_section(doc, "作用机制（MoA）", data.get("moa"))

    # timeline table
    tl = [e for e in (data.get("timeline") or []) if e.get("date")]
    if tl:
        tl.sort(key=lambda e: date_key(e.get("date")))
        doc.add_heading("研发 / 审批时间线（美国 + 中国）", level=1)
        t = doc.add_table(rows=1, cols=4)
        t.style = "Table Grid"
        for i, h in enumerate(["地区", "日期", "里程碑 / 适应症", "关键试验 + 徽标"]):
            t.rows[0].cells[i].text = ""
            rp = t.rows[0].cells[i].paragraphs[0].add_run(h)
            rp.bold = True
            shade(t.rows[0].cells[i], "EEF2F7")
        for e in tl:
            us = is_us(e.get("region"))
            cells = t.add_row().cells
            cells[0].text = ("\U0001F1FA\U0001F1F8 美国" if us else "\U0001F1E8\U0001F1F3 中国")
            shade(cells[0], "DBEAFE" if us else "FEE2E2")
            cells[1].text = str(e.get("date", ""))
            cells[2].text = str(e.get("label", ""))
            cells[3].text = str(e.get("trial", "") or "")
            add_badges(cells[3].paragraphs[0], e.get("badge"))
        set_widths(t, [1.1, 1.0, 2.7, 1.9])

    # indication matrix
    inds = data.get("indications")
    if inds:
        doc.add_heading("适应症矩阵（按地区）", level=1)
        t = doc.add_table(rows=1, cols=3)
        t.style = "Table Grid"
        for i, h in enumerate(["适应症", "\U0001F1FA\U0001F1F8 美国", "\U0001F1E8\U0001F1F3 中国"]):
            t.rows[0].cells[i].text = ""
            rp = t.rows[0].cells[i].paragraphs[0].add_run(h)
            rp.bold = True
            shade(t.rows[0].cells[i], "EEF2F7")
        for r in inds:
            cells = t.add_row().cells
            cells[0].text = ""
            rn = cells[0].paragraphs[0].add_run(str(r.get("indication", "")))
            rn.bold = True
            shade(cells[0], "F1F5F9")
            for idx, key in ((1, "us"), (2, "china")):
                c = r.get(key) or {}
                cells[idx].text = str(c.get("status", "—"))
                if c.get("trial"):
                    sp = cells[idx].add_paragraph(str(c.get("trial")))
                    sp.runs[0].font.size = Pt(8.5)
                    sp.runs[0].font.color.rgb = RGBColor.from_string("6B7280")
                add_badges(cells[idx].paragraphs[0], c.get("badge"))
        set_widths(t, [2.5, 2.0, 2.0])

    kv_section(doc, "简介", data.get("intro"))

    # sales table
    sales = data.get("sales")
    if sales and sales.get("points"):
        doc.add_heading("销售趋势  🔴 财报口径 · 每个数字都须人工复核", level=1)
        if sales.get("unit"):
            up = doc.add_paragraph("单位：" + str(sales["unit"]) + "（forecast = 预测值，非实际入账）")
            up.runs[0].font.size = Pt(9)
        rows = [["年份", "销售额", "备注 / 🔴"]]
        for p in sales["points"]:
            note = p.get("note") or ""
            if p.get("forecast"):
                note = (note + " · 预测/forecast").strip(" ·")
            rows.append([str(p.get("year", "")), str(p.get("value", "")), (note or "🔴 须复核")])
        header_table(doc, rows, [1.6, 2.2, 2.7])
        if sales.get("source"):
            sp = doc.add_paragraph("来源：" + str(sales["source"]))
            sp.runs[0].font.size = Pt(8.5)
            sp.runs[0].font.color.rgb = RGBColor.from_string("6B7280")

    kv_section(doc, "专利 / 独占到期（LOE）🔴", data.get("loe"))
    kv_section(doc, "仿制药 / 生物类似药上市？", data.get("generics"))
    kv_section(doc, "同类药品", data.get("peers"))
    kv_section(doc, "核心优势（AI 总结 · 🔴 跨试验比较不严谨）", data.get("advantages"))

    # red checklist
    ck = data.get("checklist") or []
    doc.add_heading("🔴 需人工核对清单（本工具最关键产出）", level=1)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    for i, h in enumerate(["#", "字段", "当前值 / 状态", "为什么要核", "建议去哪个官方源核"]):
        t.rows[0].cells[i].text = ""
        rp = t.rows[0].cells[i].paragraphs[0].add_run(h)
        rp.bold = True
        rp.font.color.rgb = RGBColor.from_string("991B1B")
        shade(t.rows[0].cells[i], "FEE2E2")
    for i, it in enumerate(ck, 1):
        cells = t.add_row().cells
        cells[0].text = str(i)
        cells[1].text = str(it.get("field", ""))
        cells[2].text = str(it.get("value", ""))
        add_badges(cells[2].paragraphs[0], it.get("badge"))
        cells[3].text = str(it.get("why", ""))
        cells[4].text = str(it.get("where", ""))
    set_widths(t, [0.4, 1.7, 1.7, 1.7, 1.7])

    # references
    refs = data.get("references")
    if refs:
        doc.add_heading("参考来源（Reference list）", level=1)
        doc.add_paragraph("下列为本报告所有引用链接的汇总；请优先以 🏛️官方 来源核对客观事实。")
        for r in refs:
            if isinstance(r, str):
                url, label, badge = r, r, None
            else:
                url = r.get("url") or ""
                label = r.get("label") or url
                badge = r.get("badge")
            p = doc.add_paragraph(style="List Number")
            p.add_run(str(label) + "  ")
            if url:
                add_hyperlink(p, url, url)
            add_badges(p, badge)

    out = args.out or (args.data.rsplit(".", 1)[0] + ".docx")
    doc.save(out)
    print("Wrote " + out)
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
